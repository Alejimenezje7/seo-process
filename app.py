import streamlit as st
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ──────────────────────────────────────────────
# Page config
# ──────────────────────────────────────────────
st.set_page_config(
    page_title="Mapa de Procesos SEO & ASO",
    page_icon="🗺️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ──────────────────────────────────────────────
# Custom CSS
# ──────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Rubik:wght@300;400;500;600;700&display=swap');

    html, body, [class*="css"] { font-family: 'Rubik', sans-serif; }

    /* Hide default streamlit elements */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* Card grid */
    .card-grid { display: grid; grid-template-columns: repeat(3, 1fr); gap: 24px; margin: 30px 0; }
    @media (max-width: 900px) { .card-grid { grid-template-columns: repeat(2, 1fr); } }
    @media (max-width: 640px) { .card-grid { grid-template-columns: 1fr; } }

    .card {
        background-color: #1a1a1a; border: 1px solid #333; border-radius: 8px;
        padding: 32px 24px; transition: all 0.3s ease; cursor: pointer; text-decoration: none; display: block;
    }
    .card:hover { border-color: #555; background-color: #222; transform: translateY(-4px); box-shadow: 0 8px 25px rgba(0,0,0,0.3); }
    .card-icon { font-size: 2.2rem; margin-bottom: 12px; display: block; }
    .card-title { font-size: 1.2rem; font-weight: 600; color: #ffffff; margin-bottom: 10px; }
    .card-desc { font-size: 0.9rem; color: #999; line-height: 1.6; }

    /* Section headers */
    .section-header { text-align: center; margin: 50px 0 10px; }
    .section-title { font-size: 2.2rem; font-weight: 700; color: #ffffff; letter-spacing: -1px; margin-bottom: 6px; }
    .section-subtitle { font-size: 1rem; color: #888; font-weight: 300; text-transform: uppercase; letter-spacing: 2px; }

    /* Password screen */
    .pw-container { display: flex; justify-content: center; align-items: center; min-height: 80vh; }
    .pw-card { max-width: 420px; width: 100%; background: #1a1a1a; border: 1px solid #333; border-radius: 8px; padding: 36px; }
    .pw-title { font-size: 2rem; font-weight: 700; color: #fff; margin-bottom: 8px; }
    .pw-sub { color: #999; margin-bottom: 24px; }

    /* Detail content styling */
    .detail-content h2 { color: #ffffff; margin-top: 20px; }
    .detail-content h3 { color: #dddddd; }
    .detail-content p { color: #cccccc; line-height: 1.8; }
    .detail-content li { color: #cccccc; line-height: 1.6; margin-bottom: 8px; }
    .detail-content strong { color: #ffffff; }

    /* Process flow */
    .process-flow {
        background-color: #0a0a0a; padding: 20px; border-radius: 5px;
        font-family: monospace; overflow-x: auto; margin: 20px 0;
        color: #88d498; font-size: 0.85rem; line-height: 1.6;
        white-space: pre-wrap; word-wrap: break-word; border: 1px solid #333;
    }

    /* Code block */
    .code-block {
        background-color: #0a0a0a; padding: 15px; border-radius: 5px;
        font-family: monospace; overflow-x: auto; margin: 15px 0;
        color: #88d498; font-size: 0.9rem; line-height: 1.5; border: 1px solid #333;
    }

    /* Back button */
    .back-btn {
        display: inline-flex; align-items: center; gap: 8px; color: #999;
        text-decoration: none; font-size: 0.95rem; cursor: pointer;
        background: none; border: none; padding: 8px 0; margin-bottom: 20px;
        font-family: 'Rubik', sans-serif;
    }
    .back-btn:hover { color: #fff; }

    /* Banner SVG container */
    .banner-container {
        width: 100%; overflow-x: auto; background: #1a1a1a;
        border: 1px solid #333; border-radius: 8px; margin-bottom: 30px; padding: 10px;
    }

    /* Streamlit expander overrides */
    .streamlit-expanderHeader { font-weight: 600 !important; font-size: 1.05rem !important; }

    /* Link style */
    a { color: #4a9eff; }
    a:hover { color: #6bb3ff; }
</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────
# Session state init
# ──────────────────────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "current_page" not in st.session_state:
    st.session_state.current_page = "index"


def navigate(page):
    st.session_state.current_page = page


# ──────────────────────────────────────────────
# Word document generator
# ──────────────────────────────────────────────
def generate_word(title, sections):
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for section_title, content_blocks in sections:
        doc.add_heading(section_title, level=1)
        for block in content_blocks:
            if block["type"] == "paragraph":
                doc.add_paragraph(block["text"])
            elif block["type"] == "heading2":
                doc.add_heading(block["text"], level=2)
            elif block["type"] == "heading3":
                doc.add_heading(block["text"], level=3)
            elif block["type"] == "bullet":
                doc.add_paragraph(block["text"], style="List Bullet")
            elif block["type"] == "code":
                p = doc.add_paragraph()
                run = p.add_run(block["text"])
                run.font.name = "Consolas"
                run.font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ──────────────────────────────────────────────
# PASSWORD SCREEN
# ──────────────────────────────────────────────
def render_password():
    st.markdown('<div class="pw-container"><div class="pw-card">', unsafe_allow_html=True)
    st.markdown('<div class="pw-title">Acceso privado</div>', unsafe_allow_html=True)
    st.markdown('<div class="pw-sub">Ingresa la contraseña para ver el mapa de procesos.</div>', unsafe_allow_html=True)

    password = st.text_input("Contraseña", type="password", key="pw_input")

    if st.button("Entrar", type="primary", use_container_width=True):
        if password.strip() == "SEOLAM2026":
            st.session_state.authenticated = True
            st.rerun()
        else:
            st.error("Contraseña incorrecta. Inténtalo de nuevo.")

    st.markdown('</div></div>', unsafe_allow_html=True)


# ──────────────────────────────────────────────
# INDEX PAGE
# ──────────────────────────────────────────────
def render_index():
    # SEO Section
    st.markdown("""
    <div class="section-header">
        <div class="section-title">Mapa de Procesos SEO</div>
        <div class="section-subtitle">Technical SEO</div>
    </div>
    """, unsafe_allow_html=True)

    seo_cards = [
        ("🔍", "Redirecciones 404", "Gestión de errores 404, implementación de redirecciones 301 y estrategias para preservar autoridad SEO.", "auditoria"),
        ("🔗", "Estructura de URLs", "Optimización de URLs amigables, jerarquía lógica y canonicalización de páginas.", "urls"),
        ("📱", "Mobile Optimization", "Responsive design, velocidad en móvil y optimización de Core Web Vitals.", "mobile"),
        ("🚀", "Rendimiento", "Mejora de velocidad, compresión de imágenes, caching y optimización de recursos.", "rendimiento"),
        ("🏷️", "Metadata & Schema", "Optimización de meta tags, structured data y schema markup para mejor indexación.", "metadata"),
        ("📑", "Indexación", "Gestión de sitemaps, robots.txt, crawl budget y estrategias de indexación.", "indexacion"),
    ]

    cols = st.columns(3)
    for i, (icon, title, desc, page_id) in enumerate(seo_cards):
        with cols[i % 3]:
            with st.container(border=True):
                st.markdown(f"### {icon} {title}")
                st.caption(desc)
                if st.button("Ver proceso →", key=f"btn_{page_id}", use_container_width=True):
                    navigate(page_id)
                    st.rerun()

    st.markdown("---")

    # ASO Section
    st.markdown("""
    <div class="section-header">
        <div class="section-title">Mapa de Procesos ASO</div>
        <div class="section-subtitle">ASO</div>
    </div>
    """, unsafe_allow_html=True)

    aso_cards = [
        ("⚙️", "ASO Optimizations", "Optimización de metadatos de aplicaciones, keywords, icono y descripción para mejor visibilidad en tiendas.", "aso_optimizations"),
        ("📊", "In App Events", "Configuración y seguimiento de eventos dentro de la aplicación para analytics y conversión.", "in_app_events"),
        ("🧪", "AB Test", "Diseño, ejecución y análisis de pruebas A/B para optimizar la experiencia del usuario.", "ab_test"),
    ]

    cols = st.columns(3)
    for i, (icon, title, desc, page_id) in enumerate(aso_cards):
        with cols[i % 3]:
            with st.container(border=True):
                st.markdown(f"### {icon} {title}")
                st.caption(desc)
                if st.button("Ver proceso →", key=f"btn_{page_id}", use_container_width=True):
                    navigate(page_id)
                    st.rerun()


# ──────────────────────────────────────────────
# BACK BUTTON HELPER
# ──────────────────────────────────────────────
def back_button():
    if st.button("← Volver al mapa", key="back"):
        navigate("index")
        st.rerun()


# ──────────────────────────────────────────────
# PAGE: REDIRECCIONES 404
# ──────────────────────────────────────────────
def render_auditoria():
    back_button()
    st.title("🔍 Redirecciones 404")

    # SVG Banner
    st.markdown("""
    <div class="banner-container">
    <svg viewBox="0 0 2200 350" xmlns="http://www.w3.org/2000/svg" preserveAspectRatio="xMidYMid meet" style="width:100%;min-width:1200px;">
        <rect width="2200" height="350" fill="#0a0a0a"/>
        <defs>
            <marker id="ah" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><polygon points="0 0,10 3,0 6" fill="#888"/></marker>
            <marker id="ahr" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><polygon points="0 0,10 3,0 6" fill="#ff6b6b"/></marker>
            <marker id="ahg" markerWidth="10" markerHeight="10" refX="9" refY="3" orient="auto"><polygon points="0 0,10 3,0 6" fill="#4ade80"/></marker>
        </defs>
        <circle cx="50" cy="175" r="25" fill="#1a1a1a" stroke="#4a9eff" stroke-width="2"/><text x="50" y="180" font-family="Rubik,sans-serif" font-size="12" font-weight="600" fill="#4a9eff" text-anchor="middle">Inicio</text>
        <line x1="75" y1="175" x2="115" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="115" y="155" width="120" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="175" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Identificar 404</text>
        <line x1="235" y1="175" x2="275" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="275" y="155" width="130" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="340" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Revisar GSC</text>
        <line x1="405" y1="175" x2="445" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <polygon points="495,175 545,145 595,175 545,205" fill="#1a1a1a" stroke="#ffcc00" stroke-width="2"/><text x="545" y="180" font-family="Rubik,sans-serif" font-size="10" font-weight="600" fill="#ffcc00" text-anchor="middle">¿Válido?</text>
        <line x1="545" y1="205" x2="545" y2="260" stroke="#ff6b6b" stroke-width="2" marker-end="url(#ahr)"/><text x="555" y="235" font-family="Rubik,sans-serif" font-size="10" fill="#ff6b6b" font-weight="600">No</text>
        <rect x="460" y="260" width="170" height="35" rx="4" fill="#1a1a1a" stroke="#ff6b6b" stroke-width="2"/><text x="545" y="283" font-family="Rubik,sans-serif" font-size="10" fill="#ff6b6b" text-anchor="middle">Descartar</text>
        <line x1="595" y1="175" x2="635" y2="175" stroke="#4ade80" stroke-width="2" marker-end="url(#ahg)"/><text x="610" y="165" font-family="Rubik,sans-serif" font-size="10" fill="#4ade80" font-weight="600">Sí</text>
        <rect x="635" y="155" width="120" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="695" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Crear Ticket</text>
        <line x1="755" y1="175" x2="795" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="795" y="155" width="120" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="855" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Definir Tipo</text>
        <line x1="915" y1="175" x2="955" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="955" y="155" width="110" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="1010" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Enviar</text>
        <line x1="1065" y1="175" x2="1105" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="1105" y="155" width="120" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="1165" y="177" font-family="Rubik,sans-serif" font-size="10" fill="#e0e0e0" text-anchor="middle">Esperar</text><text x="1165" y="191" font-family="Rubik,sans-serif" font-size="8" fill="#888" text-anchor="middle">(1-2 días)</text>
        <line x1="1225" y1="175" x2="1265" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="1265" y="155" width="130" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="1330" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Implementar</text>
        <line x1="1395" y1="175" x2="1435" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="1435" y="155" width="110" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="1490" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Validar</text>
        <line x1="1545" y1="175" x2="1585" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <rect x="1585" y="155" width="110" height="40" rx="4" fill="#1a1a1a" stroke="#888" stroke-width="2"/><text x="1640" y="182" font-family="Rubik,sans-serif" font-size="11" fill="#e0e0e0" text-anchor="middle">Cerrar</text>
        <line x1="1695" y1="175" x2="1735" y2="175" stroke="#888" stroke-width="2" marker-end="url(#ah)"/>
        <circle cx="1785" cy="175" r="25" fill="#1a1a1a" stroke="#4ade80" stroke-width="2"/><text x="1785" y="180" font-family="Rubik,sans-serif" font-size="12" font-weight="600" fill="#4ade80" text-anchor="middle">Fin</text>
    </svg>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("📋 Información General 404", expanded=True):
        st.markdown("""
**¿Qué es un Error 404?**

Un error 404 (Not Found) ocurre cuando un usuario intenta acceder a una página que no existe en tu servidor. Esto puede suceder por varias razones: URLs rotas, contenido eliminado, cambios en la estructura del sitio, o enlaces externos desactualizados.

**¿Por Qué las Redirecciones 404 Importan?**
- **Experiencia del Usuario:** Un 404 sin redirección crea frustración en los visitantes.
- **Pérdida de Autoridad SEO:** Los enlaces rotos disipan la autoridad de página (link juice).
- **Rastreo Ineficiente:** Google rastrea más URLs 404 que productivas, desperdiciando crawl budget.
- **Impacto en Conversiones:** Los usuarios abandonan sitios con demasiados errores 404.
- **Señal de Mantenimiento:** Los errores indican falta de cuidado y actualización del sitio.

**Tipos de Redirecciones**
- **Redirección 301 (Movido Permanentemente):** Transfiere el 90-99% de la autoridad de página.
- **Redirección 302 (Movido Temporalmente):** Se usa para cambios temporales.
- **Redirección 307 (Temporal):** Similar a 302 pero más estricto con el método HTTP.
- **Redirección Meta Refresh:** Se ejecuta en el navegador. No es recomendada para SEO.
- **Redirección JavaScript:** Implementada con código JS. Menos ideal para SEO.
        """)

    with st.expander("🔄 Proceso de Redirección"):
        st.markdown("### Mapa de Procesos — Redirección 404 en adidas")
        st.markdown("**Objetivo:** Gestionar correctamente solicitudes de redirecciones 404 para asegurar una buena experiencia de usuario y minimizar impactos SEO.")

        st.markdown("""
<div class="process-flow">Inicio
   ↓
Identificación de URL 404
   ↓
Validar si la redirección tiene sentido SEO
   ↓
Revisar información en Google Search Console
   ↓
¿La redirección es válida?
   ├── No → Descartar solicitud / reevaluar destino
   └── Sí
          ↓
Ir al template/formato de Confluence
          ↓
Completar ticket de redirección
          ↓
Definir tipo de redirección:
   ├── Directa
   └── Dinámica
          ↓
Enviar ticket
          ↓
Esperar respuesta del equipo responsable
(1–2 días hábiles aprox.)
          ↓
Implementación de la redirección
          ↓
Validar funcionamiento:
   ├── Status code correcto
   ├── URL destino correcta
   ├── Sin loops ni errores
   └── Validación SEO básica
          ↓
Agregar comentario de confirmación en el ticket
          ↓
Cerrar proceso
          ↓
Fin</div>
        """, unsafe_allow_html=True)

        st.markdown("**URL importante:**")
        st.markdown("[CDN Service Catalog - Confluence](https://confluence.tools.3stripes.net/spaces/CDN/pages/1613803193/CDN+Service+Catalog)")
        st.caption("Dependiendo del objetivo que quieras, ya sea eliminar una redirección actual o crear una nueva deberás ver su opción correspondiente al hacer scroll down")

        st.markdown("""
**1. Validación de Redirección**
- Revisar Google Search Console
- Analizar tráfico e intención de búsqueda
- Verificar similitud entre URL origen y destino
- Evitar redirecciones irrelevantes o masivas al homepage

**2. Creación del Ticket** — Herramienta: Confluence

**3. Definición del Tipo de Redirección**
""")
        st.markdown('<div class="code-block">/producto-viejo → /producto-nuevo</div>', unsafe_allow_html=True)
        st.markdown('<div class="code-block">/outlet/* → /sale/*</div>', unsafe_allow_html=True)
        st.markdown("""
**4. Espera de Implementación** — 1–2 días hábiles

**5. QA y Validación Final**
- Redirección 301 correcta, destino funcional, sin cadenas ni loops
- Herramientas: Redirect Path, DevTools, Screaming Frog, GSC

**6. Cierre del Ticket** — Confirmar, comentar, marcar completado
        """)

    with st.expander("👥 Contactos o Stakeholders"):
        st.warning("En caso URGENTE, contactar por TEAMs:")
        st.markdown("""
- **Geison Garzón:** Geison.Garzon@externals.adidas.com
- **Oksana Tretiak:** Oksana.Tretiak@externals.adidas.com
        """)

    word_data = generate_word("Redirecciones 404", [
        ("Información General", [
            {"type": "paragraph", "text": "Un error 404 ocurre cuando un usuario intenta acceder a una página que no existe. Gestionar correctamente estos errores es crucial para SEO."},
            {"type": "bullet", "text": "Experiencia del Usuario: Un 404 sin redirección crea frustración."},
            {"type": "bullet", "text": "Pérdida de Autoridad SEO: Los enlaces rotos disipan link juice."},
            {"type": "bullet", "text": "Rastreo Ineficiente: Google desperdicia crawl budget."},
        ]),
        ("Proceso de Redirección", [
            {"type": "paragraph", "text": "Objetivo: Gestionar correctamente solicitudes de redirecciones 404."},
            {"type": "heading3", "text": "1. Validación de Redirección"},
            {"type": "bullet", "text": "Revisar Google Search Console"},
            {"type": "bullet", "text": "Analizar tráfico e intención de búsqueda"},
            {"type": "heading3", "text": "2. Creación del Ticket en Confluence"},
            {"type": "heading3", "text": "3. Definición del Tipo de Redirección"},
            {"type": "code", "text": "/producto-viejo → /producto-nuevo\n/outlet/* → /sale/*"},
            {"type": "heading3", "text": "4. Espera de Implementación (1-2 días)"},
            {"type": "heading3", "text": "5. QA y Validación Final"},
            {"type": "heading3", "text": "6. Cierre del Ticket"},
        ]),
        ("Contactos", [
            {"type": "bullet", "text": "Geison Garzón: Geison.Garzon@externals.adidas.com"},
            {"type": "bullet", "text": "Oksana Tretiak: Oksana.Tretiak@externals.adidas.com"},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "Redirecciones_404.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: ESTRUCTURA DE URLs
# ──────────────────────────────────────────────
def render_urls():
    back_button()
    st.title("🔗 Estructura de URLs")

    with st.expander("📋 Información General", expanded=True):
        st.markdown("""
**Importancia de una Estructura de URLs Óptima**

La estructura de las URLs es un elemento fundamental del SEO técnico. Las URLs bien estructuradas no solo ayudan a los motores de búsqueda a entender la jerarquía y contenido de tu sitio, sino que también mejoran la experiencia del usuario.

**Principios de una Estructura URL Ideal**
- **Claridad:** Las URLs deben ser descriptivas y fáciles de entender.
- **Brevedad:** Evita URLs demasiado largas; idealmente menos de 75 caracteres.
- **Palabras Clave:** Incluye términos relevantes que describan el contenido.
- **Guiones en lugar de guiones bajos:** Utiliza guiones (-) para separar palabras.
- **Minúsculas:** Mantén todas las URLs en minúsculas.
- **Sin parámetros innecesarios:** Minimiza parámetros de consulta.
- **Estructura lógica:** Organiza las URLs siguiendo la jerarquía del sitio.

**Ejemplos**
- ✅ ejemplo.com/blog/seo-tecnico-guia-completa
- ✅ ejemplo.com/servicios/consultoria-seo
- ❌ ejemplo.com/p?id=123&cat=seo
- ❌ ejemplo.com/blog_post_seo_tecnico

**Canonicalización de URLs**

Implementa etiquetas canónicas para evitar problemas de contenido duplicado. Esta etiqueta le indica a los buscadores cuál es la versión preferida de una página.
        """)

    word_data = generate_word("Estructura de URLs", [
        ("Principios", [
            {"type": "bullet", "text": "Claridad: URLs descriptivas y fáciles de entender."},
            {"type": "bullet", "text": "Brevedad: Menos de 75 caracteres."},
            {"type": "bullet", "text": "Palabras Clave: Términos relevantes."},
            {"type": "bullet", "text": "Guiones (-) en lugar de guiones bajos (_)."},
            {"type": "bullet", "text": "Minúsculas siempre."},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "Estructura_URLs.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: MOBILE OPTIMIZATION
# ──────────────────────────────────────────────
def render_mobile():
    back_button()
    st.title("📱 Mobile Optimization")

    with st.expander("📋 Mobile-First en SEO", expanded=True):
        st.markdown("""
Google utiliza indexación **Mobile-First**, lo que significa que el rastreador de Google ve y indexa principalmente la versión móvil de tu sitio.

**Core Web Vitals**
- **LCP (Largest Contentful Paint):** Velocidad de carga → Objetivo: < 2.5s
- **FID (First Input Delay):** Interactividad → Objetivo: < 100ms
- **CLS (Cumulative Layout Shift):** Estabilidad visual → Objetivo: < 0.1

**Checklist de Optimización Móvil**
- Implementar diseño responsive para todos los tamaños de pantalla
- Asegurar que el texto sea legible sin necesidad de zoom
- Optimizar imágenes para conexiones lentas
- Evitar intersticiales intrusivos
- Usar botones con tamaño adecuado para tocar
- Minimizar redirecciones y código innecesario
- Probar en dispositivos reales y emuladores

**Testing Móvil:** Usa Google Mobile-Friendly Test, Lighthouse y Chrome DevTools.
        """)

    word_data = generate_word("Mobile Optimization", [
        ("Core Web Vitals", [
            {"type": "bullet", "text": "LCP: Velocidad de carga → < 2.5s"},
            {"type": "bullet", "text": "FID: Interactividad → < 100ms"},
            {"type": "bullet", "text": "CLS: Estabilidad visual → < 0.1"},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "Mobile_Optimization.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: RENDIMIENTO
# ──────────────────────────────────────────────
def render_rendimiento():
    back_button()
    st.title("🚀 Rendimiento")

    with st.expander("📋 Velocidad como Factor de Ranking", expanded=True):
        st.markdown("""
La velocidad de un sitio web es un factor directo de ranking en Google. Un sitio lento incrementa la tasa de rebote y reduce las conversiones.

**Estrategias de Optimización**
- **Compresión de Imágenes:** Formatos como WebP.
- **Minificación:** Reduce CSS, JavaScript y HTML.
- **Lazy Loading:** Carga contenido solo cuando es visible.
- **Caching:** Cachés a nivel de navegador y servidor.
- **CDN:** Distribuye contenido desde servidores cercanos.
- **Code Splitting:** Divide JavaScript en fragmentos.
- **Eliminación de Render-Blocking:** Optimiza recursos críticos.

**Herramientas de Monitoreo**
- Google PageSpeed Insights
- Google Lighthouse
- WebPageTest
- GTmetrix
- Pingdom

**Impacto:** Cada segundo de delay puede disminuir las conversiones entre 4-7%.
        """)

    word_data = generate_word("Rendimiento", [
        ("Estrategias", [
            {"type": "bullet", "text": "Compresión de Imágenes (WebP)"},
            {"type": "bullet", "text": "Minificación de CSS/JS/HTML"},
            {"type": "bullet", "text": "Lazy Loading"},
            {"type": "bullet", "text": "Caching"},
            {"type": "bullet", "text": "CDN"},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "Rendimiento.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: METADATA & SCHEMA
# ──────────────────────────────────────────────
def render_metadata():
    back_button()
    st.title("🏷️ Metadata & Schema")

    with st.expander("📋 Metadata: El Puente con los Buscadores", expanded=True):
        st.markdown("""
Los meta tags son etiquetas HTML que proporcionan información sobre el contenido de tu página. Son cruciales para el SEO.

**Meta Tags Esenciales**
- **Title Tag:** Máximo 60-70 caracteres. Incluir palabras clave principales.
- **Meta Description:** 150-160 caracteres. Atractivo para mejorar CTR.
- **Meta Keywords:** Menos importante ahora, pero buena práctica.
- **Meta Robots:** Controla rastreo e indexación.
- **Open Graph Tags:** Optimiza compartir en redes sociales.
- **Twitter Card Tags:** Específicas para Twitter.

**Schema Markup (Datos Estructurados)**

Ayudan a los motores de búsqueda a entender mejor el contenido. Puede resultar en rich snippets.
- Schema Organization
- Schema Product
- Schema Article
- Schema FAQ
- Schema LocalBusiness
- Schema Review

**JSON-LD vs Microdata:** JSON-LD es el formato recomendado por Google. Más fácil de implementar y mantener.
        """)

    word_data = generate_word("Metadata & Schema", [
        ("Meta Tags Esenciales", [
            {"type": "bullet", "text": "Title Tag: 60-70 caracteres con keywords"},
            {"type": "bullet", "text": "Meta Description: 150-160 caracteres"},
            {"type": "bullet", "text": "Meta Robots, Open Graph, Twitter Cards"},
        ]),
        ("Schema Markup", [
            {"type": "bullet", "text": "Organization, Product, Article, FAQ, LocalBusiness, Review"},
            {"type": "paragraph", "text": "JSON-LD es el formato recomendado por Google."},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "Metadata_Schema.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: INDEXACIÓN
# ──────────────────────────────────────────────
def render_indexacion():
    back_button()
    st.title("📑 Indexación")

    with st.expander("📋 ¿Qué es la Indexación?", expanded=True):
        st.markdown("""
La indexación es el proceso mediante el cual los motores de búsqueda descubren, rastrean y añaden páginas a su base de datos.

**Archivo Robots.txt**
- Define qué carpetas o archivos deben ser evitados
- Especifica la ubicación del sitemap XML
- Controla la velocidad de rastreo (crawl delay)
- Previene indexación de contenido duplicado o privado

**Sitemap XML**
- Incluye todas las URLs principales del sitio
- Proporciona metadatos como fecha de última modificación
- Define la prioridad relativa de páginas
- Declara frequency de actualización sugerida

**Gestión de Crawl Budget**
- Elimina páginas duplicadas o de bajo valor
- Bloquea parámetros de sesión innecesarios
- Utiliza redirects 301 en lugar de content duplicado
- Monitorea el rastreo en Google Search Console

**Exclusión Inteligente:** No todas las páginas deben ser indexadas. Excluir páginas innecesarias mejora la eficiencia.
        """)

    word_data = generate_word("Indexación", [
        ("Robots.txt", [{"type": "bullet", "text": "Define carpetas a evitar, ubicación del sitemap, crawl delay"}]),
        ("Sitemap XML", [{"type": "bullet", "text": "URLs principales, metadatos, prioridad, frequency"}]),
        ("Crawl Budget", [{"type": "bullet", "text": "Eliminar duplicados, bloquear parámetros, usar 301s, monitorear GSC"}]),
    ])
    st.download_button("📥 Descargar Word", word_data, "Indexacion.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: ASO OPTIMIZATIONS
# ──────────────────────────────────────────────
def render_aso_optimizations():
    back_button()
    st.title("⚙️ ASO Optimizations")

    with st.expander("📋 Información General ASO", expanded=True):
        st.markdown("""
**¿Qué es ASO (App Store Optimization)?**

ASO es el proceso de optimización de aplicaciones móviles para mejorar su visibilidad en las tiendas de aplicaciones (App Store, Google Play).

**¿Por Qué ASO es Importante?**
- **Mayor Visibilidad:** Mejora el ranking en búsquedas.
- **Aumento de Descargas:** Más descargas sin costo de adquisición.
- **Mejor Conversión:** Un listing optimizado convierte más visitantes.
- **Competitividad:** En mercados saturados, ASO marca la diferencia.
- **Retención:** Atraer usuarios relevantes mejora la retención.

**Elementos Clave de ASO**
- **Nombre de la Aplicación:** Palabras clave relevantes y memorable.
- **Keywords:** Alto volumen y baja competencia.
- **Descripción:** Clara, concisa y orientada a beneficios.
- **Icono:** Visualmente atractivo y distinguible.
- **Capturas de Pantalla:** Funciones principales y generar interés.
- **Calificaciones y Reseñas:** Influyen en descarga y algoritmo.
        """)

    with st.expander("🔄 Proceso de Optimización"):
        st.markdown("""
**1. Investigación de Keywords**
- Analizar palabras clave de competidores
- Usar App Annie, Sensor Tower o Mobile Action
- Identificar keywords de alto volumen y baja competencia

**2. Optimización de Metadatos**
- Nombre con keyword principal, subtitle, descripción corta

**3. Diseño Visual**
- Icono, capturas de pantalla, preview video

**4. Configuración en Tiendas**
- Metadatos, categoría, precio, compatibilidad

**5. Testing y Validación**
- Búsquedas con keywords, listing correcto, monitorear ranking

**6. Monitoreo y Ajustes**
- Rankings, conversión, reseñas, ajustes según datos
        """)

    with st.expander("👥 Contactos o Stakeholders"):
        st.markdown("""
- **Ana López** - ASO Manager
- **David Pérez** - App Marketing Specialist
- **Elena Rodríguez** - UX/UI Designer
        """)

    word_data = generate_word("ASO Optimizations", [
        ("Elementos Clave", [
            {"type": "bullet", "text": "Nombre, Keywords, Descripción, Icono, Capturas, Reseñas"},
        ]),
        ("Proceso", [
            {"type": "bullet", "text": "1. Investigación de Keywords → 2. Metadatos → 3. Diseño Visual → 4. Config Tiendas → 5. Testing → 6. Monitoreo"},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "ASO_Optimizations.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: IN APP EVENTS
# ──────────────────────────────────────────────
def render_in_app_events():
    back_button()
    st.title("📊 In App Events")

    with st.expander("📋 Información General In App Events", expanded=True):
        st.markdown("""
**¿Qué son In App Events?**

Los In App Events son acciones específicas que los usuarios realizan dentro de una aplicación móvil y que son rastreadas mediante herramientas de analytics. Generan datos cruciales para entender el comportamiento del usuario, optimizar la experiencia y medir el éxito de campañas.

**¿Por Qué Son Importantes?**
- **Medición de Conversiones:** Rastrean acciones clave que llevan a conversión.
- **Análisis del Comportamiento:** Insight sobre cómo interactúan los usuarios.
- **Optimización:** Identificar y optimizar puntos débiles.
- **Retargeting:** Campañas de remarketing dirigidas.
- **ROI:** Demostrar valor de la inversión en marketing móvil.

**Tipos de In App Events**
- **Custom Events:** Personalizados según necesidades del negocio.
- **Standard Events:** Predefinidos como Purchase, AddToCart, Signup.
- **Funnel Events:** Secuencia que conforma un flujo de conversión.
- **Revenue Events:** Relacionados con monetización e ingresos.
        """)

    with st.expander("🔄 Proceso de Configuración"):
        st.markdown("""
**Objetivo:** Organizar, planificar y activar In App Events para LAM de forma coordinada con Global ASO y Content, asegurando disponibilidad de slots, fechas claras y materiales listos antes del go live.

**1. Validación de slots con Global ASO**

Tener llamadas recurrentes con **Federica Bello** de Global ASO para conocer los slots disponibles. Generalmente, al ser LAM pionero en la implementación, tenemos varios slots disponibles.

**2. Cronograma con Content**

Reunión con **Alexa Peralta** de Content. Documento disponible en:
""")
        st.markdown("[📄 LAM - ASO - COMMERCIAR EVENTS - Organizer](https://adidasgroup-my.sharepoint.com/:x:/r/personal/jose_rubertone_adidas_com/_layouts/15/Doc.aspx?sourcedoc=%7BBD2C4A94-E8A9-42EB-8093-87C9DB56043F%7D&file=LAM%20-%20ASO%20-%20COMMERCIAR%20EVENTS%20-%20Organizer.xlsx&action=default&mobileredirect=true)")
        st.markdown("""
Colocar próximos In App Events con fecha de go live y fecha de terminación.

**3. Push a Content por país**

Seguimiento con la persona de Content responsable por país. Materiales idealmente **7 días antes** del go live.
        """)

    with st.expander("🛠️ Configuración en la Herramienta"):
        st.markdown("""
**Checklist inicial**
- ✅ Validar que el slot esté confirmado con Global ASO
- ✅ Confirmar fecha de go live y terminación en el cronograma
- ✅ Revisar que Content haya entregado materiales finales
- ✅ Verificar textos, imágenes, país, fechas y objetivo

**Materiales necesarios**
- Nombre del In App Event
- Descripción o copy aprobado
- Assets visuales finales
- País o países donde se activará
- Fecha y hora de inicio
- Fecha y hora de finalización

**Validación final**
- Revisar la vista previa antes de enviar
- Confirmar que las fechas coincidan con el cronograma
- Compartir confirmación con ASO y Content una vez configurado
        """)

    with st.expander("👥 Contactos o Stakeholders"):
        st.markdown("""
- **Federica Bello** — Global ASO. Confirma disponibilidad de slots.
- **Alexa Peralta** — Content. Apoya construcción del cronograma.
- **Content por país** — Entrega materiales 7 días antes del go live.
        """)

    word_data = generate_word("In App Events", [
        ("Tipos de Events", [
            {"type": "bullet", "text": "Custom Events, Standard Events, Funnel Events, Revenue Events"},
        ]),
        ("Proceso de Configuración", [
            {"type": "bullet", "text": "1. Validación de slots con Global ASO (Federica Bello)"},
            {"type": "bullet", "text": "2. Cronograma con Content (Alexa Peralta)"},
            {"type": "bullet", "text": "3. Push a Content por país — materiales 7 días antes del go live"},
        ]),
        ("Contactos", [
            {"type": "bullet", "text": "Federica Bello — Global ASO"},
            {"type": "bullet", "text": "Alexa Peralta — Content"},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "In_App_Events.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# PAGE: AB TEST
# ──────────────────────────────────────────────
def render_ab_test():
    back_button()
    st.title("🧪 AB Test")

    with st.expander("📋 Información General AB Testing", expanded=True):
        st.markdown("""
**¿Qué es AB Testing?**

Metodología que compara dos versiones de una página o elemento para determinar cuál performa mejor. División aleatoria del tráfico entre dos variantes para tomar decisiones informadas.

**¿Por Qué es Importante?**
- **Decisiones Basadas en Datos:** Elimina la especulación.
- **Mejora Continua:** Iterar y optimizar constantemente.
- **Reducción de Riesgo:** Valida cambios antes de implementación completa.
- **Incremento de Conversiones:** Mejoras acumulativas generan resultados significativos.
- **ROI Óptimo:** Asegura que los cambios generan valor.

**Elementos Comunes para AB Testing**
- **Copy y Messaging:** Títulos, botones, descripciones.
- **Design:** Colores, layout, tipografía, espaciado.
- **Elementos Visuales:** Imágenes, iconos, animaciones.
- **User Flow:** Orden de pasos, ubicación de botones.
- **Precios y Ofertas:** Valores, planes, descuentos.
        """)

    with st.expander("🔄 Proceso de AB Testing"):
        st.markdown("""
**Objetivo:** Ejecutar pruebas A/B rigurosas para validar cambios de diseño, mensajería y funcionalidad.

**1. Identificación de Hipótesis**
- Analizar métricas actuales e identificar oportunidades
- Plantear hipótesis clara: "Si cambio X, entonces Z mejorará"
- Priorizar por impacto potencial y facilidad
- Documentar baseline actual

**2. Diseño del Test**
- Versión A (Control): Sin cambios
- Versión B (Variante): Con el cambio a probar
- Métrica clave, tamaño de muestra, duración (mín. 2 semanas)

**3. Configuración Técnica**
- Plataforma: Optimizely, Apptimize, Firebase
- Segmentación y audiencia target
- Distribución 50/50

**4. Ejecución y Monitoreo**
- Monitorear métricas diariamente
- Detectar anomalías — NO detener prematuramente

**5. Análisis de Resultados**
- Diferencia entre A y B
- Significancia estadística (p-value < 0.05)
- Documentar aprendizajes

**6. Decisión e Implementación**
- **Si B gana:** Implementar al 100%
- **Si empate:** Evaluar otros factores
- **Si A gana:** Mantener control o probar variante diferente
        """)

    with st.expander("👥 Contactos o Stakeholders"):
        st.markdown("""
- **María Flores** - Product Manager
- **Roberto Jiménez** - UX Researcher
- **Camila Soto** - Growth Analyst
        """)

    word_data = generate_word("AB Test", [
        ("Proceso", [
            {"type": "bullet", "text": "1. Identificación de Hipótesis"},
            {"type": "bullet", "text": "2. Diseño del Test (Control vs Variante)"},
            {"type": "bullet", "text": "3. Configuración Técnica (Optimizely/Firebase)"},
            {"type": "bullet", "text": "4. Ejecución y Monitoreo"},
            {"type": "bullet", "text": "5. Análisis de Resultados (p-value < 0.05)"},
            {"type": "bullet", "text": "6. Decisión e Implementación"},
        ]),
    ])
    st.download_button("📥 Descargar Word", word_data, "AB_Test.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ──────────────────────────────────────────────
# ROUTER
# ──────────────────────────────────────────────
PAGE_MAP = {
    "index": render_index,
    "auditoria": render_auditoria,
    "urls": render_urls,
    "mobile": render_mobile,
    "rendimiento": render_rendimiento,
    "metadata": render_metadata,
    "indexacion": render_indexacion,
    "aso_optimizations": render_aso_optimizations,
    "in_app_events": render_in_app_events,
    "ab_test": render_ab_test,
}

if not st.session_state.authenticated:
    render_password()
else:
    # Sidebar navigation
    with st.sidebar:
        st.markdown("### 🗺️ Navegación")
        st.markdown("---")
        st.markdown("**SEO**")
        if st.button("🏠 Inicio", use_container_width=True, key="nav_home"):
            navigate("index"); st.rerun()
        if st.button("🔍 Redirecciones 404", use_container_width=True, key="nav_404"):
            navigate("auditoria"); st.rerun()
        if st.button("🔗 Estructura de URLs", use_container_width=True, key="nav_urls"):
            navigate("urls"); st.rerun()
        if st.button("📱 Mobile Optimization", use_container_width=True, key="nav_mobile"):
            navigate("mobile"); st.rerun()
        if st.button("🚀 Rendimiento", use_container_width=True, key="nav_rend"):
            navigate("rendimiento"); st.rerun()
        if st.button("🏷️ Metadata & Schema", use_container_width=True, key="nav_meta"):
            navigate("metadata"); st.rerun()
        if st.button("📑 Indexación", use_container_width=True, key="nav_idx"):
            navigate("indexacion"); st.rerun()

        st.markdown("---")
        st.markdown("**ASO**")
        if st.button("⚙️ ASO Optimizations", use_container_width=True, key="nav_aso"):
            navigate("aso_optimizations"); st.rerun()
        if st.button("📊 In App Events", use_container_width=True, key="nav_events"):
            navigate("in_app_events"); st.rerun()
        if st.button("🧪 AB Test", use_container_width=True, key="nav_ab"):
            navigate("ab_test"); st.rerun()

        st.markdown("---")
        if st.button("🚪 Cerrar sesión", use_container_width=True, key="nav_logout"):
            st.session_state.authenticated = False
            st.session_state.current_page = "index"
            st.rerun()

    renderer = PAGE_MAP.get(st.session_state.current_page, render_index)
    renderer()
